"""Author the proposal into the branded Scale Computing Word template and convert
it to PDF.

The template (resources/TMPL - Generic Document Template_2025.docx) carries the
branding — first-page watermark/title header, running header/footer, logo, and
named styles (Title, Heading 1/2, Normal). We strip its instructional body
content but keep the section properties (so headers/footers/watermark survive),
then author the proposal using those styles. The deck content is mirrored but
"leads with the recommendation".

PDF is produced by converting the authored .docx with headless LibreOffice
(soffice), which is installed in the container.
"""

import io
import os
import shutil
import subprocess
import tempfile
import threading

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

from export_pptx import _svg_to_png_bytes, _fmt_ram, _fmt_num
from export_gauges import render_util_bars, util_rows, compute_floor_sentence
from recommend import _rec_network_svg
from cluster_diagram import render_replication_topology_svg
from i18n import translator, font_for, is_cjk

_TEMPLATE = os.path.join(os.path.dirname(__file__), "..", "resources",
                         "TMPL - Generic Document Template_2025.docx")

DK2 = RGBColor(0x11, 0x38, 0x59)
MUTED = RGBColor(0x5A, 0x6B, 0x7D)

# The template's named styles carry the Martel Sans branding font. Martel can't
# render CJK scripts, so for those languages we force a CJK-capable font on the
# text runs (Latin is left untouched so branded layout stays pixel-identical).
BODY_FONT = "Martel Sans"


def _apply_lang_font(run, lang):
    """For CJK languages, set an explicit CJK-capable font on the run (both the
    Latin ascii/hAnsi slot and the eastAsia slot) so text renders instead of
    tofu. For Latin languages this is a no-op — runs keep inheriting the
    template style's Martel Sans variant, preserving the branded typography."""
    if not is_cjk(lang):
        return
    name = font_for(lang, BODY_FONT)
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), name)


def _apply_lang_font_para(par, lang):
    """Apply the CJK font to every run in an already-built paragraph (used for
    heading/title/bullet paragraphs authored via add_heading/add_paragraph)."""
    if not is_cjk(lang):
        return
    for run in par.runs:
        _apply_lang_font(run, lang)


def _clear_body(doc):
    """Remove the template's instructional body content but keep the final
    sectPr (which references the branded headers/footers and page setup)."""
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _set_header_text(section, text):
    """Replace the template's 'Header' placeholder in both the first-page and
    running headers, preserving its (grey Martel Sans) styling. The placeholder
    lives inside a text box (vertical-text shape), so it isn't reachable via
    paragraphs/runs — walk every w:t in the header part instead (catches both the
    shape and its mc:Fallback copy)."""
    for hdr in (section.first_page_header, section.header):
        if hdr is None:
            continue
        for t in hdr._element.iter(qn("w:t")):
            if (t.text or "").strip() == "Header":
                t.text = text


def _content_width(doc):
    """Usable text width (inches) from the template's actual page + margins, so
    tables match the paragraph block exactly (template uses 0.75" margins → 7.0")."""
    sec = doc.sections[0]
    return (sec.page_width - sec.left_margin - sec.right_margin) / 914400


# OOXML requires tblPr children in this exact order. Word ENFORCES it (and
# silently drops misordered tblW/tblLayout → falls back to autofit → tables
# overflow the page); LibreOffice is lenient, which is why the PDF looked fine
# while the .docx in Word did not. We must insert in-order, not append.
_TBLPR_ORDER = [
    "w:tblStyle", "w:tblpPr", "w:tblOverlap", "w:bidiVisual",
    "w:tblStyleRowBandSize", "w:tblStyleColBandSize", "w:tblW", "w:jc",
    "w:tblCellSpacing", "w:tblInd", "w:tblBorders", "w:shd", "w:tblLayout",
    "w:tblCellMar", "w:tblLook", "w:tblCaption", "w:tblDescription",
]


def _set_tblpr_child(tblPr, tag):
    """Replace (or create) a tblPr child, inserted at its schema-correct position."""
    for el in tblPr.findall(qn(tag)):
        tblPr.remove(el)
    el = OxmlElement(tag)
    successors = {qn(t) for t in _TBLPR_ORDER[_TBLPR_ORDER.index(tag) + 1:]}
    for child in tblPr:
        if child.tag in successors:
            child.addprevious(el)
            return el
    tblPr.append(el)
    return el


def _shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _set_table_borders(table, color="DDE2E6"):
    borders = _set_tblpr_child(table._tbl.tblPr, "w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color)
        borders.append(el)


def _cell_margins(table, top=70, bottom=70, left=130, right=130):
    """Breathing room inside every cell (twips)."""
    mar = _set_tblpr_child(table._tbl.tblPr, "w:tblCellMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        mar.append(e)


def _fixed_layout(table, widths):
    """Lock column widths so long content WRAPS instead of overflowing the page.
    Sets tblW (preferred width), tblLayout=fixed, tblInd=0, the authoritative
    tblGrid columns, and each cell's tcW — all inserted in schema order so Word
    honours them."""
    table.autofit = False
    table.allow_autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr

    tblW = _set_tblpr_child(tblPr, "w:tblW")
    tblW.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tblW.set(qn("w:type"), "dxa")

    tblInd = _set_tblpr_child(tblPr, "w:tblInd")
    tblInd.set(qn("w:w"), "0")
    tblInd.set(qn("w:type"), "dxa")

    layout = _set_tblpr_child(tblPr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")

    grid_cols = tbl.tblGrid.findall(qn("w:gridCol"))
    for i, w in enumerate(widths):
        if i < len(grid_cols):
            grid_cols[i].set(qn("w:w"), str(int(w * 1440)))
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w)


def _para_keep_next(p_el):
    """Set <w:keepNext/> on a raw <w:p> element, in schema-correct position
    (after pStyle). Idempotent."""
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_el.insert(0, pPr)
    if pPr.findall(qn("w:keepNext")):
        return
    kn = OxmlElement("w:keepNext")
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is not None:
        pStyle.addnext(kn)
    else:
        pPr.insert(0, kn)


def _pin_heading_to_table(table):
    """Keep a table's heading on the same page as the table, so a heading (and its
    short intro paragraph) is never orphaned at the foot of a page when the table
    is pushed to the next. Walks the table's preceding sibling paragraphs marking
    each 'keep with next', and stops once it pins the section heading (or hits a
    non-paragraph — e.g. an earlier table — or a small step cap, so it never
    chains back into the previous section). keep-with-next is a soft constraint:
    Word still breaks if the heading+intro+table genuinely can't fit one page."""
    el = table._tbl.getprevious()
    steps = 0
    while el is not None and el.tag == qn("w:p") and steps < 5:
        pPr = el.find(qn("w:pPr"))
        style = ""
        if pPr is not None:
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is not None:
                style = (pStyle.get(qn("w:val")) or "").lower()
        is_heading = style.startswith("heading") or style.startswith("title")
        _para_keep_next(el)
        if is_heading:
            break  # reached the section heading — done
        el = el.getprevious()
        steps += 1


def _keep_table_together(table):
    """Keep a table from being split across pages: mark every row 'cannot split'
    (no row breaks mid-cell) and 'keep with next' on all rows but the last, so
    Word holds the whole table on one page and pushes it to the next page when it
    won't fit. Also pins the heading above it (see _pin_heading_to_table). A
    table taller than a single page still breaks — Word overrides keep-with-next
    once the content exceeds the page, which is the desired behaviour."""
    rows = table.rows
    last = len(rows) - 1
    for ri, row in enumerate(rows):
        trPr = row._tr.get_or_add_trPr()
        if not trPr.findall(qn("w:cantSplit")):
            trPr.append(OxmlElement("w:cantSplit"))
        if ri < last:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.keep_with_next = True
    _pin_heading_to_table(table)


def _style_cell(cell, text, bold=False, color=None, fill=None, align=None, lang="en"):
    cell.text = ""
    pr = cell.paragraphs[0]
    if align is not None:
        pr.alignment = align
    run = pr.add_run(str(text))
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    _apply_lang_font(run, lang)
    if fill is not None:
        _shade(cell, fill)


def _spec_table(doc, rows, total_w, label_w=2.5, lang="en"):
    """Clean two-column spec table: shaded bold labels (left) + values (right),
    spanning the full content width so its right edge matches the paragraphs."""
    t = doc.add_table(rows=0, cols=2)
    _set_table_borders(t)
    _cell_margins(t)
    for k, v in rows:
        cells = t.add_row().cells
        _style_cell(cells[0], k, bold=True, color=DK2, fill="EEF2F7", lang=lang)
        _style_cell(cells[1], v, lang=lang)
    _fixed_layout(t, [label_w, total_w - label_w])
    _keep_table_together(t)
    return t


def _grid_table(doc, headers, rows, total_w, weights=None, lang="en"):
    """Header-row (dark) + data-rows table spanning the full content width.
    weights are relative column proportions (default equal)."""
    t = doc.add_table(rows=1, cols=len(headers))
    _set_table_borders(t)
    _cell_margins(t)
    for i, h in enumerate(headers):
        _style_cell(t.rows[0].cells[i], h, bold=True,
                    color=RGBColor(0xFF, 0xFF, 0xFF), fill="113859", lang=lang)
    for r in rows:
        cells = t.add_row().cells
        for i, val in enumerate(r):
            _style_cell(cells[i], val, lang=lang)
    if weights is None:
        weights = [1] * len(headers)
    scale = total_w / sum(weights)
    _fixed_layout(t, [w * scale for w in weights])
    _keep_table_together(t)
    return t


def _spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)
    p.paragraph_format.space_before = Pt(0)
    return p


def _para(doc, text, italic=False, color=None, size=None, lang="en"):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)
    _apply_lang_font(run, lang)
    return p


def _finalize_doc(doc):
    # Final pass: clear any indent inherited from docDefaults on every paragraph.
    for par in doc.paragraphs:
        ppf = par.paragraph_format
        ppf.left_indent = Inches(0)
        ppf.right_indent = Inches(0)
        ppf.first_line_indent = Inches(0)


def _doc_setup(doc, lang, header_text):
    """Shared document-level setup: comfortable margins, zeroed style indents,
    and the running header text."""
    # The template ships with tight 0.75" (~1.9 cm) side margins, which leaves the
    # body running to the page edge. Widen to a comfortable 1" (2.54 cm) so there's
    # real whitespace on the right and full-width tables sit inside the page.
    sec = doc.sections[0]
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    # The template's paragraph styles inherit a left indent, so the body sits inset
    # from the title/tables (a visible "double" left margin). Zero it on every style
    # we use so text aligns to the page margin; a final per-paragraph pass before
    # save catches anything inherited from docDefaults.
    for sname in ("Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        try:
            spf = doc.styles[sname].paragraph_format
            spf.left_indent = Inches(0)
            spf.right_indent = Inches(0)
            spf.first_line_indent = Inches(0)
        except KeyError:
            pass
    _set_header_text(sec, header_text)
    return sec


def build_proposal_docx(summary, recommendation, projection, source_perf=None, lang="en"):
    t9n = translator(lang)
    doc = Document(_TEMPLATE) if os.path.exists(_TEMPLATE) else Document()
    _clear_body(doc)
    # Header: proposal title plus the customer/cluster name when one is set.
    _hdr_name = (summary.get("cluster_name") or "").strip()
    _doc_setup(doc, lang, t9n("export.docx.title")
               + (f" — {_hdr_name}" if _hdr_name else ""))

    # ── Title ────────────────────────────────────────────────────────────────
    _apply_lang_font_para(doc.add_paragraph(t9n("export.docx.title"), style="Title"), lang)
    subtitle = summary.get("cluster_name") or summary.get("current_platform") or ""
    if subtitle:
        _apply_lang_font_para(doc.add_paragraph(subtitle, style="Subtitle"), lang)

    _append_proposal_body(doc, summary, recommendation, projection, source_perf, lang)
    _finalize_doc(doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _heading(doc, text, lang, level=1):
    h = doc.add_heading(text, level=level)
    _apply_lang_font_para(h, lang)
    return h


def _bullet(doc, text, lang):
    style = "List Bullet" if "List Bullet" in [st.name for st in doc.styles] else "Normal"
    b = doc.add_paragraph(text, style=style)
    _apply_lang_font_para(b, lang)
    return b


def _is_dedicated(summary):
    """A dedicated DR target has no primary workload of its own."""
    return (summary.get("host_count", 0) == 0 and summary.get("active_vms", 0) == 0)


def _platform_capabilities(doc, lang):
    """The Scale Computing HyperCore platform pitch — management, HEAT, SCRIBE,
    AIME. Emitted ONCE in a multi-site proposal (not repeated per site)."""
    t9n = translator(lang)
    _heading(doc, t9n("export.docx.multisite_platform_heading"), lang, level=1)
    _para(doc, t9n("export.docx.product_intro"), lang=lang)

    _heading(doc, t9n("export.docx.mgmt_operations"), lang, level=2)
    _para(doc, t9n("export.docx.mgmt_operations_intro"), lang=lang)
    for k in ("export.docx.mgmt_bullet_1", "export.docx.mgmt_bullet_2",
              "export.docx.mgmt_bullet_3", "export.docx.mgmt_bullet_4",
              "export.docx.mgmt_bullet_5"):
        _bullet(doc, t9n(k), lang)

    _heading(doc, t9n("export.docx.heat_heading"), lang, level=2)
    _para(doc, t9n("export.docx.heat_intro"), lang=lang)
    for k in ("export.docx.heat_bullet_1", "export.docx.heat_bullet_2",
              "export.docx.heat_bullet_3", "export.docx.heat_bullet_4",
              "export.docx.heat_bullet_5"):
        _bullet(doc, t9n(k), lang)

    _heading(doc, t9n("export.docx.scribe_heading"), lang, level=2)
    _para(doc, t9n("export.docx.scribe_intro"), lang=lang)
    for k in ("export.docx.scribe_bullet_1", "export.docx.scribe_bullet_2",
              "export.docx.scribe_bullet_3", "export.docx.scribe_bullet_4",
              "export.docx.scribe_bullet_5", "export.docx.scribe_bullet_6",
              "export.docx.scribe_bullet_7"):
        _bullet(doc, t9n(k), lang)

    _heading(doc, t9n("export.docx.aime_heading"), lang, level=2)
    _para(doc, t9n("export.docx.aime_intro"), lang=lang)
    for k in ("export.docx.aime_bullet_1", "export.docx.aime_bullet_2",
              "export.docx.aime_bullet_3", "export.docx.aime_bullet_4",
              "export.docx.aime_bullet_5", "export.docx.aime_bullet_6",
              "export.docx.aime_bullet_7"):
        _bullet(doc, t9n(k), lang)


def _append_site_sizing(doc, cl, lang, cw):
    """Per-site engineering detail only (config, network, sizing rationale,
    capacity, benchmark) — no repeated marketing. Used by the multi-site doc."""
    t9n = translator(lang)
    r = cl["recommendation"]
    s = cl["summary"]
    p = cl["projection"]
    t = r["totals"]
    n1 = r["n_minus_1"]
    iops = r.get("iops") or {}
    dedicated = _is_dedicated(s)

    so = r.get("storage_only")
    hci_nodes = r.get("hci_node_count", r["node_count"])
    nodes_label = (t9n("export.common.hci_plus_storage_only", hci=hci_nodes, storage=so["count"])
                   if so else t9n("export.common.node_count", count=r["node_count"]))
    num_cl = r.get("num_clusters", 1)
    cl_label = (t9n("export.common.clusters_layout", count=num_cl,
                    layout=" + ".join(map(str, r.get("cluster_layout", []))))
                if num_cl > 1 else t9n("export.common.single_cluster"))

    # Site intro — dedicated DR vs a workload-bearing site.
    if dedicated:
        _para(doc, t9n("export.docx.multisite_dedicated_intro", name=cl.get("name", "")), lang=lang)
    else:
        _para(doc, t9n("export.docx.multisite_site_intro",
                       name=cl.get("name", ""), vms=s.get("active_vms", 0),
                       vcpus=_fmt_num(s.get("total_vcpus", 0)),
                       ram=_fmt_ram(s.get("total_vm_provisioned_memory_gb", 0)),
                       used_tb=s.get("datastore_used_tb", 0), years=p["years"]), lang=lang)

    # Recommended configuration
    _heading(doc, t9n("export.docx.recommended_configuration"), lang, level=2)
    spec_rows = [
        (t9n("export.docx.recommended_platform"), f"{r['model']} · {nodes_label} · {cl_label}"),
        (t9n("export.docx.per_node_cpu"), r["cpu"]),
        (t9n("export.docx.per_node_cores_threads"),
         t9n("export.docx.cores_threads_val", cores=r["cores_per_node"], threads=r["threads_per_node"])),
        (t9n("export.docx.per_node_ram"), _fmt_ram(r["ram_per_node_gb"])),
        (t9n("export.docx.per_node_storage"), r["storage_config"]["desc"]),
        (t9n("export.docx.cluster_cores"), str(t["cores"])),
        (t9n("export.docx.cluster_ram"), _fmt_ram(t["ram_gb"])),
        (t9n("export.docx.cluster_usable_storage"), f"{t['usable_storage_tb']} TB"),
    ]
    if iops:
        spec_rows.append((t9n("export.docx.cluster_net_iops"), f"{iops['total']:,}"))
    spec_rows += [
        (t9n("export.docx.n1_resilient"),
         t9n("export.docx.n1_resilient_val", cores=n1["cores"],
             ram=_fmt_ram(n1["ram_gb"]), usable_tb=n1["usable_storage_tb"])),
        (t9n("export.docx.vcpu_core_ratio"), f"{r['vcpu_ratio']:.2f} : 1"),
    ]
    _spec_table(doc, spec_rows, total_w=cw, lang=lang)
    if r.get("single_node"):
        _para(doc, t9n("export.common.single_node_note"), italic=True, color=MUTED, lang=lang)
    _spacer(doc)

    # Network diagram
    svg = _rec_network_svg(r, lang) or r.get("network_svg")
    if svg:
        png = _svg_to_png_bytes(svg, out_width=2200)
        if png:
            _heading(doc, t9n("export.docx.cluster_network"), lang, level=3)
            doc.add_picture(io.BytesIO(png), width=Inches(min(6.5, cw)))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sizing rationale (utilization bars + determinant)
    u = r.get("utilization")
    if u:
        rows_u, any_ha = util_rows(u)
        if rows_u:
            _heading(doc, t9n("export.docx.sizing_rationale"), lang, level=2)
            png = render_util_bars(rows_u, limiting_key=(r.get("determinant") or {}).get("resource", ""),
                                   any_ha=any_ha, lang=lang)
            doc.add_picture(io.BytesIO(png), width=Inches(cw))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _spacer(doc, 4)
            det = r.get("determinant") or {}
            res, hr = det.get("resource"), det.get("headroom_pct")
            if res == "CPU":
                _para(doc, t9n("export.docx.rationale_cpu", vcpus=s.get("total_vcpus"),
                               ratio=f"{r.get('vcpu_ratio', 0):.2f}",
                               required=f"{det.get('required'):.0f}", achieved=f"{det.get('achieved'):.0f}",
                               headroom=f"{hr:.1f}"), lang=lang)
            elif res in ("RAM", "Storage"):
                _para(doc, t9n("export.docx.rationale_ram_storage", resource=res,
                               required=det.get("required"), unit=det.get("unit", ""),
                               achieved=det.get("achieved"), headroom=f"{hr:.1f}"), lang=lang)
            elif res == "Compute":
                cf = r.get("compute_floor") or {}
                _para(doc, t9n("export.docx.rationale_compute", achieved=f"{det.get('achieved'):.0f}",
                               util=f"{cf.get('source_cpu_util_pct', 100):.0f}"), lang=lang)
            _para(doc, t9n("export.docx.rationale_bar_legend"), italic=True, size=9, lang=lang)
            _spacer(doc)

    # Capacity planning (skip for a dedicated DR — its growth mirrors the sources)
    if not dedicated:
        _heading(doc, t9n("export.docx.capacity_planning", years=p["years"]), lang, level=2)
        _grid_table(doc,
                    [t9n("export.common.resource"), t9n("export.common.current"),
                     t9n("export.docx.year_n", years=p["years"]), t9n("export.docx.proposed_n1")],
                    [["vCPUs", _fmt_num(p["base_vcpus"]), _fmt_num(p["projected_vcpus"]),
                      t9n("export.docx.cores_at_ratio", cores=n1["cores"], ratio=f"{r['vcpu_ratio']:.1f}")],
                     ["RAM", _fmt_ram(p["base_ram_gb"]), _fmt_ram(p["projected_ram_gb"]), _fmt_ram(n1["ram_gb"])],
                     [t9n("export.common.storage"), f"{p['base_storage_tb']} TB",
                      f"{p['projected_storage_tb']} TB", t9n("export.docx.tb_usable", tb=n1["usable_storage_tb"])]],
                    total_w=cw, weights=[1.6, 1.8, 1.8, 1.8], lang=lang)
        _spacer(doc)

    # Benchmark vs current (only when a source benchmark was provided)
    tgt = t.get("perf_index")
    sp = cl.get("source_perf")
    if sp and sp.get("total_specrate") and tgt:
        _heading(doc, t9n("export.docx.performance_vs_current"), lang, level=2)
        src_total = sp["total_specrate"]
        ratio = tgt / src_total if src_total else 0
        verdict = (t9n("export.docx.verdict_multiple", ratio=f"{ratio:.1f}") if ratio >= 1
                   else t9n("export.docx.verdict_fraction", pct=round(ratio * 100)))
        _spec_table(doc, [
            (t9n("export.docx.recommended_cluster"), f"{r.get('cpu', '')} × {r.get('node_count', '')} nodes"),
            (t9n("export.docx.cluster_specrate2017"), _fmt_num(tgt)),
            (t9n("export.docx.benchmark_should_deliver"), verdict),
        ], total_w=cw, label_w=2.6, lang=lang)
        _spacer(doc)


def build_multisite_proposal_docx(clusters, lang="en"):
    """A single commercial multi-site proposal with one narrative flow:
    executive summary → solution at a glance (overview + topology) → why
    HyperCore (platform pitch, once) → per-site sizing detail → assumptions.
    `clusters` = [{name, summary, recommendation, projection, source_perf,
    replicates_to}]."""
    t9n = translator(lang)
    doc = Document(_TEMPLATE) if os.path.exists(_TEMPLATE) else Document()
    _clear_body(doc)
    _doc_setup(doc, lang, t9n("export.docx.multisite_title"))
    cw = _content_width(doc)

    _apply_lang_font_para(doc.add_paragraph(t9n("export.docx.multisite_title"), style="Title"), lang)
    _apply_lang_font_para(doc.add_paragraph(
        t9n("export.docx.multisite_subtitle", count=len(clusters)), style="Subtitle"), lang)

    # ── aggregate figures across all sites ────────────────────────────────────
    sites = len(clusters)
    agg_hosts = sum(cl["summary"].get("host_count", 0) for cl in clusters)
    agg_vms = sum(cl["summary"].get("active_vms", 0) for cl in clusters)
    agg_used_tb = round(sum(cl["summary"].get("datastore_used_tb", 0) for cl in clusters), 1)
    tot_nodes = sum(cl["recommendation"].get("node_count", 0) for cl in clusters)
    tot_cores = sum(cl["recommendation"]["totals"].get("cores", 0) for cl in clusters)
    tot_ram = sum(cl["recommendation"]["totals"].get("ram_gb", 0) for cl in clusters)
    tot_usable = round(sum(cl["recommendation"]["totals"].get("usable_storage_tb", 0) for cl in clusters), 1)
    years = clusters[0]["projection"].get("years", 5) if clusters else 5

    # ── Executive summary ─────────────────────────────────────────────────────
    _heading(doc, t9n("export.docx.multisite_exec_heading"), lang, level=1)
    _para(doc, t9n("export.docx.multisite_exec_intro",
                   sites=sites, hosts=agg_hosts, vms=agg_vms, used_tb=agg_used_tb,
                   years=years, nodes=tot_nodes, cores=_fmt_num(tot_cores),
                   ram=_fmt_ram(tot_ram), usable_tb=tot_usable), lang=lang)

    # ── Solution at a glance: overview table + combined totals + topology ─────
    _heading(doc, t9n("export.docx.multisite_solution_heading"), lang, level=1)
    show_rep = any(cl.get("replicates_to") for cl in clusters)
    header = [t9n("export.pptx.multisite_col_cluster"), t9n("export.pptx.multisite_col_model"),
              t9n("export.pptx.multisite_col_nodes"), t9n("export.pptx.multisite_col_cores"),
              t9n("export.pptx.multisite_col_ram"), t9n("export.pptx.multisite_col_storage")]
    if show_rep:
        header.append(t9n("export.pptx.multisite_col_replicates"))
    rows = []
    for cl in clusters:
        r = cl["recommendation"]
        tot = r.get("totals", {})
        row = [cl.get("name", ""), r.get("model", ""), str(r.get("node_count", "")),
               str(tot.get("cores", "")), _fmt_ram(tot.get("ram_gb", 0)),
               f"{tot.get('usable_storage_tb', 0)} TB"]
        if show_rep:
            row.append(cl.get("replicates_to") or "—")
        rows.append(row)
    weights = [1.5, 2.0, 0.8, 0.9, 1.1, 1.2, 1.3] if show_rep else [1.6, 2.2, 0.9, 1.0, 1.2, 1.3]
    _grid_table(doc, header, rows, total_w=cw, weights=weights, lang=lang)
    _spacer(doc)
    _spec_table(doc, [
        (t9n("export.docx.multisite_total_nodes"), str(tot_nodes)),
        (t9n("export.docx.multisite_total_cores"), _fmt_num(tot_cores)),
        (t9n("export.docx.multisite_total_ram"), _fmt_ram(tot_ram)),
        (t9n("export.docx.multisite_total_storage"), f"{tot_usable} TB"),
    ], total_w=cw, label_w=2.6, lang=lang)
    _spacer(doc)

    # Replication & DR (topology diagram + one explanation) — only if configured.
    topo_svg = render_replication_topology_svg(clusters, lang)
    if topo_svg:
        png = _svg_to_png_bytes(topo_svg, out_width=2200)
        if png:
            _heading(doc, t9n("export.docx.multisite_replication_heading"), lang, level=2)
            _para(doc, t9n("export.docx.multisite_replication_intro"), lang=lang)
            doc.add_picture(io.BytesIO(png), width=Inches(min(6.5, cw)))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _spacer(doc)

            # SC//Connect — the secure connectivity layer that carries the
            # routed replication traffic between sites (cross-sell).
            _heading(doc, t9n("export.docx.sc_connect_heading"), lang, level=3)
            _para(doc, t9n("export.docx.sc_connect_intro"), lang=lang)
            for k in ("export.docx.sc_connect_bullet_1", "export.docx.sc_connect_bullet_2",
                      "export.docx.sc_connect_bullet_3", "export.docx.sc_connect_bullet_4",
                      "export.docx.sc_connect_bullet_5"):
                _bullet(doc, t9n(k), lang)
            _spacer(doc)

    # ── Why Scale Computing HyperCore (platform pitch, ONCE) ──────────────────
    _platform_capabilities(doc, lang)

    # ── Per-site sizing detail ────────────────────────────────────────────────
    for cl in clusters:
        doc.add_page_break()
        _heading(doc, t9n("export.docx.multisite_site_heading", name=cl.get("name", "")), lang, level=1)
        _append_site_sizing(doc, cl, lang, cw)

    # ── Assumptions (once) ────────────────────────────────────────────────────
    _spacer(doc)
    _heading(doc, t9n("export.docx.assumptions"), lang, level=1)
    _bullet(doc, t9n("export.docx.multisite_assumption_sites"), lang)
    _bullet(doc, t9n("export.docx.assumption_3"), lang)

    _finalize_doc(doc)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _append_proposal_body(doc, summary, recommendation, projection, source_perf=None, lang="en"):
    """Append the full proposal body (management overview → platform sections)
    for one recommendation. Shared by the single- and multi-cluster builders."""
    t9n = translator(lang)

    def _add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        _apply_lang_font_para(h, lang)
        return h

    def _add_bullet(text):
        b = doc.add_paragraph(text, style=bullet_style)
        _apply_lang_font_para(b, lang)
        return b

    r = recommendation
    s = summary
    p = projection
    t = r["totals"]
    n1 = r["n_minus_1"]
    iops = r.get("iops") or {}
    bullet_style = "List Bullet" if "List Bullet" in [st.name for st in doc.styles] else "Normal"
    cw = _content_width(doc)

    so = r.get("storage_only")
    hci_nodes = r.get("hci_node_count", r["node_count"])
    nodes_label = (t9n("export.common.hci_plus_storage_only",
                       hci=hci_nodes, storage=so["count"])
                   if so else t9n("export.common.node_count", count=r["node_count"]))
    num_cl = r.get("num_clusters", 1)
    cl_label = (t9n("export.common.clusters_layout", count=num_cl,
                    layout=" + ".join(map(str, r.get("cluster_layout", []))))
                if num_cl > 1 else t9n("export.common.single_cluster"))

    # ── Management overview (executive summary — leads the document) ──────────
    _add_heading(t9n("export.docx.mgmt_overview"), level=1)
    _para(doc,
          t9n("export.docx.mgmt_overview_intro",
              platform=s.get("current_platform", "virtualization"),
              hosts=s.get("host_count", 0), vms=s.get("active_vms", 0),
              used_tb=s.get("datastore_used_tb", 0), nodes=nodes_label,
              model=r["model"], usable_tb=t["usable_storage_tb"], cores=t["cores"],
              years=p["years"], growth=p["growth_pct"],
              ratio=f"{r['vcpu_ratio']:.2f}"),
          lang=lang)
    _para(doc, t9n("export.docx.product_intro"), lang=lang)
    _spacer(doc, 4)
    fits = (p["projected_storage_tb"] <= n1["usable_storage_tb"]
            and p["projected_vcpus"] <= n1["cores"] * r["vcpu_ratio"])
    proj_fits = (t9n("export.docx.proj_fits_within") if fits
                 else t9n("export.docx.proj_fits_approaching"))
    _spec_table(doc, [
        (t9n("export.docx.recommended_platform"), f"{r['model']} · {nodes_label} · {cl_label}"),
        (t9n("export.docx.usable_capacity"),
         t9n("export.docx.usable_capacity_val",
             usable_tb=t["usable_storage_tb"], n1_tb=n1["usable_storage_tb"])),
        (t9n("export.docx.compute"),
         t9n("export.docx.compute_val", cores=t["cores"], ratio=f"{r['vcpu_ratio']:.2f}")),
        (t9n("export.docx.year_outlook", years=p["years"]),
         t9n("export.docx.year_outlook_val",
             vcpus=_fmt_num(p["projected_vcpus"]),
             storage_tb=p["projected_storage_tb"], fits=proj_fits)),
    ], total_w=cw, label_w=2.5, lang=lang)
    _spacer(doc)

    # ── Recommended configuration ────────────────────────────────────────────
    _add_heading(t9n("export.docx.recommended_configuration"), level=1)
    _para(doc, t9n("export.docx.recommended_config_intro",
                   model=r["model"], nodes=nodes_label, clusters=cl_label,
                   form_factor=r["form_factor"], chassis=r["chassis"]),
          lang=lang)
    spec_rows = [
        (t9n("export.docx.per_node_cpu"), r["cpu"]),
        (t9n("export.docx.per_node_cores_threads"),
         t9n("export.docx.cores_threads_val",
             cores=r["cores_per_node"], threads=r["threads_per_node"])),
        (t9n("export.docx.per_node_ram"), _fmt_ram(r["ram_per_node_gb"])),
        (t9n("export.docx.per_node_storage"), r["storage_config"]["desc"]),
        (t9n("export.docx.cluster_cores"), str(t["cores"])),
        (t9n("export.docx.cluster_ram"), _fmt_ram(t["ram_gb"])),
        (t9n("export.docx.cluster_usable_storage"), f"{t['usable_storage_tb']} TB"),
    ]
    if iops:
        spec_rows.append((t9n("export.docx.cluster_net_iops"), f"{iops['total']:,}"))
    spec_rows += [
        (t9n("export.docx.n1_resilient"),
         t9n("export.docx.n1_resilient_val", cores=n1["cores"],
             ram=_fmt_ram(n1["ram_gb"]), usable_tb=n1["usable_storage_tb"])),
        (t9n("export.docx.vcpu_core_ratio"), f"{r['vcpu_ratio']:.2f} : 1"),
    ]
    _spec_table(doc, spec_rows, total_w=cw, lang=lang)
    # Single-node DR target: no failover redundancy.
    if r.get("single_node"):
        _para(doc, t9n("export.common.single_node_note"), italic=True, color=MUTED, lang=lang)
    _spacer(doc)

    # Network diagram — regenerate in the document language (fall back to stored).
    svg = _rec_network_svg(r, lang) or r.get("network_svg")
    if svg:
        png = _svg_to_png_bytes(svg, out_width=2200)
        if png:
            _add_heading(t9n("export.docx.cluster_network"), level=2)
            doc.add_picture(io.BytesIO(png), width=Inches(min(6.5, cw)))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Start the sizing rationale on a fresh page (only when the diagram
            # was actually rendered, so we never emit a stray blank page).
            doc.add_page_break()

    # ── Sizing rationale (utilization bars + how the node count was reached) ──
    u = r.get("utilization")
    if u:
        rows_u, any_ha = util_rows(u)
        if rows_u:
            _add_heading(t9n("export.docx.sizing_rationale"), level=1)
            png = render_util_bars(
                rows_u, limiting_key=(r.get("determinant") or {}).get("resource", ""),
                any_ha=any_ha, lang=lang)
            doc.add_picture(io.BytesIO(png), width=Inches(cw))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _spacer(doc, 4)
            det = r.get("determinant") or {}
            res, hr = det.get("resource"), det.get("headroom_pct")
            if res == "CPU":
                _para(doc, t9n("export.docx.rationale_cpu",
                               vcpus=s.get("total_vcpus"),
                               ratio=f"{r.get('vcpu_ratio', 0):.2f}",
                               required=f"{det.get('required'):.0f}",
                               achieved=f"{det.get('achieved'):.0f}",
                               headroom=f"{hr:.1f}"), lang=lang)
            elif res in ("RAM", "Storage"):
                unit = det.get("unit", "")
                _para(doc, t9n("export.docx.rationale_ram_storage",
                               resource=res, required=det.get("required"),
                               unit=unit, achieved=det.get("achieved"),
                               headroom=f"{hr:.1f}"), lang=lang)
            elif res == "Compute":
                cf = r.get("compute_floor") or {}
                util = cf.get("source_cpu_util_pct", 100)
                _para(doc, t9n("export.docx.rationale_compute",
                               achieved=f"{det.get('achieved'):.0f}",
                               util=f"{util:.0f}"), lang=lang)
            # Show compute-floor coverage even when another resource was binding.
            if res != "Compute":
                cfs = compute_floor_sentence(r, lang)
                if cfs:
                    _para(doc, cfs, italic=True, size=9, lang=lang)
            _para(doc, t9n("export.docx.rationale_bar_legend"),
                  italic=True, size=9, lang=lang)
            _spacer(doc)

    # ── Performance vs current environment (benchmark comparison) ─────────────
    tgt = t.get("perf_index")
    if source_perf and source_perf.get("total_specrate") and tgt:
        _add_heading(t9n("export.docx.performance_vs_current"), level=1)
        src_total = source_perf["total_specrate"]
        ratio = tgt / src_total if src_total else 0
        used_pm = False
        grid_rows = []
        for c in source_perf.get("cpus", []):
            is_pm = c.get("type") == "passmark"
            used_pm = used_pm or is_pm
            grid_rows.append([c.get("model", ""), str(c.get("sockets", "")),
                              f"{_fmt_num(c.get('score', 0))} {'PassMark' if is_pm else 'SPECrate'}",
                              _fmt_num(c.get("total", 0))])
        grid_rows.append([t9n("export.docx.total_environment"), "", "", _fmt_num(src_total)])
        _grid_table(doc, [t9n("export.docx.your_current_cpus"),
                          t9n("export.docx.sockets"), t9n("export.docx.score"),
                          "SPECrate"],
                    grid_rows, total_w=cw, weights=[3.4, 1.0, 1.6, 1.0], lang=lang)
        _spacer(doc, 4)
        used_pm = used_pm or bool(r.get("cpu_perf_is_passmark"))
        verdict = (t9n("export.docx.verdict_multiple", ratio=f"{ratio:.1f}")
                   if ratio >= 1 else
                   t9n("export.docx.verdict_fraction", pct=round(ratio * 100)))
        _spec_table(doc, [
            (t9n("export.docx.recommended_cluster"),
             f"{r.get('cpu', '')} × {r.get('node_count', '')} nodes"),
            (t9n("export.docx.cluster_specrate2017"), _fmt_num(tgt)),
            (t9n("export.docx.benchmark_should_deliver"), verdict),
        ], total_w=cw, label_w=2.6, lang=lang)
        _spacer(doc, 4)
        if used_pm:
            _para(doc, t9n("export.docx.passmark_note"), italic=True, size=9, lang=lang)
        _para(doc, t9n("export.docx.benchmark_disclaimer"),
              italic=True, size=9, lang=lang)
        _spacer(doc)

    # ── Management & operations ──────────────────────────────────────────────
    _add_heading(t9n("export.docx.mgmt_operations"), level=1)
    _para(doc, t9n("export.docx.mgmt_operations_intro"), lang=lang)
    for key in ("export.docx.mgmt_bullet_1", "export.docx.mgmt_bullet_2",
                "export.docx.mgmt_bullet_3", "export.docx.mgmt_bullet_4",
                "export.docx.mgmt_bullet_5"):
        _add_bullet(t9n(key))
    _spacer(doc)

    # ── Current environment & workload ───────────────────────────────────────
    _add_heading(t9n("export.docx.current_environment"), level=1)
    _spec_table(doc, [
        (t9n("export.common.platform"), s.get("current_platform", "")),
        (t9n("export.common.hosts"),
         t9n("export.docx.hosts_val", count=s.get("host_count", 0),
             cores=_fmt_num(s.get("total_host_cores", 0)),
             ram=_fmt_ram(s.get("total_host_ram_gb", 0)))),
        (t9n("export.common.vms"),
         t9n("export.docx.vms_val", active=s.get("active_vms", 0),
             total=s.get("total_vms", 0))),
        (t9n("export.docx.workload"),
         t9n("export.docx.workload_val",
             vcpus=_fmt_num(s.get("total_vcpus", 0)),
             ram=_fmt_ram(s.get("total_vm_provisioned_memory_gb", 0)),
             used_tb=s.get("datastore_used_tb", 0))),
        (t9n("export.docx.measured_ratio"), f"{s.get('vcpu_per_core_ratio', 0):.2f} : 1"),
    ], total_w=cw, lang=lang)
    _spacer(doc)

    # ── Capacity planning ────────────────────────────────────────────────────
    _add_heading(t9n("export.docx.capacity_planning", years=p["years"]), level=1)
    _para(doc, t9n("export.docx.capacity_planning_intro",
                   growth=p["growth_pct"], snapshot=p["snapshot_pct"],
                   factor=p.get("growth_factor", 1)),
          italic=True, color=MUTED, lang=lang)
    _grid_table(doc,
                [t9n("export.common.resource"), t9n("export.common.current"),
                 t9n("export.docx.year_n", years=p["years"]),
                 t9n("export.docx.proposed_n1")],
                [["vCPUs", _fmt_num(p["base_vcpus"]), _fmt_num(p["projected_vcpus"]),
                  t9n("export.docx.cores_at_ratio", cores=n1["cores"],
                      ratio=f"{r['vcpu_ratio']:.1f}")],
                 ["RAM", _fmt_ram(p["base_ram_gb"]), _fmt_ram(p["projected_ram_gb"]),
                  _fmt_ram(n1["ram_gb"])],
                 [t9n("export.common.storage"), f"{p['base_storage_tb']} TB",
                  f"{p['projected_storage_tb']} TB",
                  t9n("export.docx.tb_usable", tb=n1["usable_storage_tb"])]],
                total_w=cw, weights=[1.6, 1.8, 1.8, 1.8], lang=lang)
    _spacer(doc)

    # ── Assumptions ──────────────────────────────────────────────────────────
    _add_heading(t9n("export.docx.assumptions"), level=1)
    _add_bullet(t9n("export.docx.assumption_1", clusters=cl_label))
    _add_bullet(t9n("export.docx.assumption_2", ratio=f"{r['vcpu_ratio']:.2f}"))
    _add_bullet(t9n("export.docx.assumption_3"))

    # ── HEAT automated tiering ───────────────────────────────────────────────
    _add_heading(t9n("export.docx.heat_heading"), level=1)
    _para(doc, t9n("export.docx.heat_intro"), lang=lang)
    for key in ("export.docx.heat_bullet_1", "export.docx.heat_bullet_2",
                "export.docx.heat_bullet_3", "export.docx.heat_bullet_4",
                "export.docx.heat_bullet_5"):
        _add_bullet(t9n(key))

    # ── SCRIBE block engine ──────────────────────────────────────────────────
    _add_heading(t9n("export.docx.scribe_heading"), level=1)
    _para(doc, t9n("export.docx.scribe_intro"), lang=lang)
    for key in ("export.docx.scribe_bullet_1", "export.docx.scribe_bullet_2",
                "export.docx.scribe_bullet_3", "export.docx.scribe_bullet_4",
                "export.docx.scribe_bullet_5", "export.docx.scribe_bullet_6",
                "export.docx.scribe_bullet_7"):
        _add_bullet(t9n(key))

    # ── AIME autonomous infrastructure management ────────────────────────────
    _add_heading(t9n("export.docx.aime_heading"), level=1)
    _para(doc, t9n("export.docx.aime_intro"), lang=lang)
    for key in ("export.docx.aime_bullet_1", "export.docx.aime_bullet_2",
                "export.docx.aime_bullet_3", "export.docx.aime_bullet_4",
                "export.docx.aime_bullet_5", "export.docx.aime_bullet_6",
                "export.docx.aime_bullet_7"):
        _add_bullet(t9n(key))


def _soffice_bin():
    return shutil.which("soffice") or shutil.which("libreoffice")


# Bounded queue in front of LibreOffice. Each soffice conversion pegs roughly a
# full core, so an uncoordinated burst of exports would thrash a small box (and
# spike RAM). The per-client rate limits (20/min) don't bound *simultaneous*
# conversions, so we gate them here: a small number convert at once, a few more
# wait in line, and anything beyond that sheds (caller → 503) instead of piling
# onto a saturated box.
#
#   _LO_MAX        - conversions running at once (semaphore "slots")
#   _LO_QUEUE_MAX  - extra requests allowed to WAIT for a slot
#   admitted cap   - _LO_MAX + _LO_QUEUE_MAX (running + waiting)
#
# NOTE: these are per-process (a worker never sees another worker's counters),
# so the effective global figures are ~(gunicorn workers) × each. With the
# defaults — 3 workers × (1 running + 3 waiting) — that's ≤3 concurrent soffice
# on a 4-core box (one core of headroom) and ≤12 total export requests in flight
# before shedding. Each waiting request holds a gunicorn thread, so keep
# (_LO_MAX + _LO_QUEUE_MAX) below the worker thread count (6) to leave threads
# for light traffic. Raise _LO_MAX only if you also add cores.
_LO_MAX = max(1, int(os.environ.get("LIBREOFFICE_MAX_CONCURRENCY", "1")))
_LO_QUEUE_MAX = max(0, int(os.environ.get("LIBREOFFICE_QUEUE_MAX", "3")))
_lo_semaphore = threading.BoundedSemaphore(_LO_MAX)

# Admission counter (running + waiting) guarded by its own lock, so we can reject
# at the queue-full boundary before a request ever blocks on the semaphore.
_lo_admit_lock = threading.Lock()
_lo_outstanding = 0

# How long a queued conversion waits for a free slot before giving up. Kept
# comfortably under the difference between the gunicorn worker timeout (180s)
# and the soffice conversion timeout (120s) so that wait + convert can't exceed
# the worker timeout and get the worker reaped mid-export. This also stops a
# single hung conversion from holding the whole queue hostage — those behind it
# time out here and shed rather than waiting indefinitely.
_LO_ACQUIRE_TIMEOUT = 45


def _office_to_pdf(data_bytes, in_ext):
    """Convert an office document (.docx/.pptx) → PDF bytes via headless
    LibreOffice. Returns None if LibreOffice isn't available, if the conversion
    queue is full or the wait window elapsed, or on conversion failure. Callers
    already treat None as "PDF unavailable" (→ HTTP 503)."""
    global _lo_outstanding
    soffice = _soffice_bin()
    if not soffice:
        return None

    # Admission control: reject immediately if running + waiting is already at
    # the cap, so the line can't grow past _LO_QUEUE_MAX and starve the worker
    # threads that light traffic needs.
    with _lo_admit_lock:
        if _lo_outstanding >= _LO_MAX + _LO_QUEUE_MAX:
            return None
        _lo_outstanding += 1
    try:
        # Admitted — wait in line for a conversion slot (bounded by the timeout).
        if not _lo_semaphore.acquire(timeout=_LO_ACQUIRE_TIMEOUT):
            return None
        try:
            with tempfile.TemporaryDirectory() as d:
                src = os.path.join(d, f"doc.{in_ext}")
                with open(src, "wb") as f:
                    f.write(data_bytes)
                try:
                    subprocess.run([soffice, "--headless", "--convert-to", "pdf",
                                    "--outdir", d, src],
                                   check=True, timeout=120,
                                   stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
                                   env={**os.environ, "HOME": d})
                except Exception:
                    return None
                out = os.path.join(d, "doc.pdf")
                return open(out, "rb").read() if os.path.exists(out) else None
        finally:
            _lo_semaphore.release()
    finally:
        with _lo_admit_lock:
            _lo_outstanding -= 1


def convert_docx_to_pdf(docx_bytes):
    return _office_to_pdf(docx_bytes, "docx")


def convert_pptx_to_pdf(pptx_bytes):
    return _office_to_pdf(pptx_bytes, "pptx")
